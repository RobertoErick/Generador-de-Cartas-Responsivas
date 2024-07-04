import tkinter as tk
from tkinter import messagebox
from tkinter import font
from docx import Document
import os

# Función para reemplazar texto en los marcadores
def reemplazar_marcadores(doc, respuestas):
    for marcador, respuesta in respuestas.items():
        for p in doc.paragraphs:
            if f'{{{{{marcador}}}}}' in p.text:
                p.text = p.text.replace(f'{{{{{marcador}}}}}', respuesta)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if f'{{{{{marcador}}}}}' in p.text:
                            p.text = p.text.replace(f'{{{{{marcador}}}}}', respuesta)

# Función para crear los directorios si no existen
def crear_directorios(base_path):
    responsivas_dir = os.path.join(base_path, 'Cartas Responsivas')
    recepcion_dir = os.path.join(base_path, 'Cartas Recepción')
    
    os.makedirs(responsivas_dir, exist_ok=True)
    os.makedirs(recepcion_dir, exist_ok=True)
    
    return responsivas_dir, recepcion_dir

# Función para crear el documento basado en una plantilla
def crear_carta_responsiva_desde_plantilla(plantilla_path, respuestas, output_dir):
    # Cargar la plantilla
    doc = Document(plantilla_path)

    # Reemplazar los marcadores con las respuestas
    reemplazar_marcadores(doc, respuestas)

    # Formatear el nombre del archivo
    nombre_archivo = f"{respuestas['Nombre']}_{respuestas['Puesto']}_{respuestas['Equipo']}_Carta_Responsiva.docx"
    output_path = os.path.join(output_dir, nombre_archivo)

    # Guardar el documento con las respuestas reemplazadas
    doc.save(output_path)
    messagebox.showinfo("Éxito", f"Documento '{nombre_archivo}' creado con éxito en {output_path}.")

# Función para crear el documento basado en una plantilla
def crear_carta_recepcion_desde_plantilla(plantilla_path, respuestas, output_dir):
    # Cargar la plantilla
    doc = Document(plantilla_path)

    # Reemplazar los marcadores con las respuestas
    reemplazar_marcadores(doc, respuestas)

    # Formatear el nombre del archivo
    nombre_archivo = f"{respuestas['Nombre']}_{respuestas['Puesto']}_{respuestas['Equipo']}_Carta_Recepción.docx"
    output_path = os.path.join(output_dir, nombre_archivo)

    # Guardar el documento con las respuestas reemplazadas
    doc.save(output_path)
    messagebox.showinfo("Éxito", f"Documento '{nombre_archivo}' creado con éxito en {output_path}.")

# Función para manejar la lógica del cuestionario
def cuestionario():
    equipo = equipo_var.get()
    if equipo == "Seleccionar":
        messagebox.showwarning("Advertencia", "Debes seleccionar un equipo.")
        return

    respuestas = {
        "Equipo": equipo,
        "Marca": entry_marca.get(),
        "Modelo": entry_modelo.get(),
        "Pulgadas": entry_pulgadas.get(),
        "Impermeable": entry_impermeable.get(),
        "Estado": entry_estado.get(),
        "Color": entry_color.get(),
        "NoSerie": entry_no_serie.get(),
        "IMEI": entry_IMEI.get(),
        "Procesador": entry_procesador.get(),
        "RAM": entry_ram.get(),
        "SistemaOperativo": entry_sistema_operativo.get(),
        "Almacenamiento": entry_almacenamiento.get(),
        "Clave": entry_clave.get(),
        "Cargador": entry_cargador.get(),
        "Mouse": entry_mouse.get(),
        "Puesto": entry_puesto.get(),
        "Nombre": entry_nombre.get()
    }

    base_path = os.getcwd()

    if equipo == "Laptop":
        plantilla_responsiva_path = os.path.join(base_path, 'plantilla_responsiva_laptop.docx')
        plantilla_recepcion_path = os.path.join(base_path, 'plantilla_recepcion_laptop.docx')
    elif equipo == "Celular":
        plantilla_responsiva_path = os.path.join(base_path, 'plantilla_responsiva_celular.docx')
        plantilla_recepcion_path = os.path.join(base_path, 'plantilla_recepcion_celular.docx')
    elif equipo == "Mochila":
        plantilla_responsiva_path = os.path.join(base_path, 'plantilla_responsiva_mochila.docx')
        plantilla_recepcion_path = os.path.join(base_path, 'plantilla_recepcion_mochila.docx')
    else:
        messagebox.showwarning("Advertencia", "Plantillas no encontradas para el equipo seleccionado.")
        return

    responsivas_dir, recepcion_dir = crear_directorios(base_path)

    crear_carta_responsiva_desde_plantilla(plantilla_responsiva_path, respuestas, responsivas_dir)
    crear_carta_recepcion_desde_plantilla(plantilla_recepcion_path, respuestas, recepcion_dir)

# Función para actualizar los campos visibles según la selección de equipo
def actualizar_campos(*args):
    equipo_seleccionado = equipo_var.get()
    campos_laptop = ["Modelo", "No. Serie", "Procesador", "RAM", "Sistema Operativo", "Almacenamiento", "Clave", "Cargador", "Mouse", "Puesto", "Nombre"]
    campos_celular = ["Modelo", "No. Serie", "IMEI", "Procesador", "RAM", "Sistema Operativo", "Almacenamiento", "Clave", "Cargador", "Puesto", "Nombre"]
    campos_mochila = ["Marca", "Modelo", "Pulgadas", "Impermeable", "Estado", "Color", "Nombre"]

    for i, (label_text, entry_var) in enumerate(labels_entries):
        if equipo_seleccionado == "Laptop" and label_text not in campos_laptop:
            globals()[entry_var].grid_remove()
            labels[i].grid_remove()
        elif equipo_seleccionado == "Celular" and label_text not in campos_celular:
            globals()[entry_var].grid_remove()
            labels[i].grid_remove()
        elif equipo_seleccionado == "Mochila" and label_text not in campos_mochila:
            globals()[entry_var].grid_remove()
            labels[i].grid_remove()
        elif equipo_seleccionado == "Seleccionar":
            globals()[entry_var].grid_remove()
            labels[i].grid_remove()
            mensaje_seleccion.grid()
        else:
            globals()[entry_var].grid()
            labels[i].grid()
            mensaje_seleccion.grid_remove()

# Crear la interfaz gráfica
root = tk.Tk()
root.title("Generador de Cartas")
root.geometry("800x800")

# Establecer la fuente
fuente_grande = font.Font(family="Helvetica", size=11)

# Crear y posicionar los widgets
labels_entries = [
    ("Marca", "entry_marca"),
    ("Modelo", "entry_modelo"),
    ("Pulgadas", "entry_pulgadas"),
    ("Impermeable", "entry_impermeable"),
    ("Estado", "entry_estado"),
    ("Color", "entry_color"),
    ("No. Serie", "entry_no_serie"),
    ("IMEI", "entry_IMEI"),
    ("Procesador", "entry_procesador"),
    ("RAM", "entry_ram"),
    ("Sistema Operativo", "entry_sistema_operativo"),
    ("Almacenamiento", "entry_almacenamiento"),
    ("Clave", "entry_clave"),
    ("Cargador", "entry_cargador"),
    ("Mouse", "entry_mouse"),
    ("Puesto", "entry_puesto"),
    ("Nombre", "entry_nombre")
]

labels = []

# Crear y posicionar el menú desplegable para "Equipo"
tk.Label(root, text="Equipo", font=fuente_grande).grid(row=0, column=0, padx=10, pady=10)
equipo_var = tk.StringVar(root)
equipo_var.set("Seleccionar")  # Valor por defecto
equipo_var.trace('w', actualizar_campos)
equipo_menu = tk.OptionMenu(root, equipo_var, "Laptop", "Celular", "Mochila")
equipo_menu.config(font=fuente_grande)
equipo_menu.grid(row=0, column=1, padx=10, pady=10)

mensaje_seleccion = tk.Label(root, text="Selecciona algún equipo", font=fuente_grande)
mensaje_seleccion.grid(row=1, columnspan=2, pady=10)

for i, (label_text, entry_var) in enumerate(labels_entries, start=2):
    label = tk.Label(root, text=label_text, font=fuente_grande)
    label.grid(row=i, column=0, padx=10, pady=10)
    labels.append(label)
    entry = tk.Entry(root, font=fuente_grande)
    entry.grid(row=i, column=1, padx=10, pady=10)
    globals()[entry_var] = entry
    label.grid_remove()
    entry.grid_remove()

submit_button = tk.Button(root, text="Generar Cartas", command=cuestionario, font=fuente_grande)
submit_button.grid(row=len(labels_entries)+2, columnspan=2, pady=20)

# Ejecutar la aplicación
root.mainloop()
